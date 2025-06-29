# Required libraries:
# pip3 install python-docx dashscope nltk Pillow tqdm
import os
import io
import re
import sys
from copy import deepcopy
from typing import Union 
import subprocess

import tempfile
import concurrent.futures
import nltk

from docx import Document
from docx.shared import Inches, Emu, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn 
from docx.text.paragraph import Paragraph as DocxParagraphObject
from docx.table import Table as DocxTableObject, _Row as DocxRowObject, _Cell as DocxCellObject

# --- Define Namespace URIs and a specific map for VML XPath calls ---
VML_NAMESPACE_URI = 'urn:schemas-microsoft-com:vml'
VML_NS_MAP = {'v': VML_NAMESPACE_URI}


# --- NLTK Punkt Tokenizer Check ---
import nltk
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    print("NLTK 'punkt' tokenizer not found. Attempting to download...")
    try:
        nltk.download('punkt', quiet=True)
        nltk.data.find('tokenizers/punkt')
        print("NLTK 'punkt' tokenizer downloaded and now available.")
    except Exception as e:
        print(f"Error downloading 'punkt': {e}")
        print("--------------------------------------------------------------------")
        print("请尝试在Python解释器中手动下载 'punkt' 模块。")
        print("--------------------------------------------------------------------")
        raise RuntimeError("NLTK 'punkt' tokenizer is required but could not be downloaded.") from e

import dashscope
from http import HTTPStatus
from tqdm import tqdm

# --- Configuration ---
QIANWEN_API_KEY = "sk-912a379c69be4924b042ec6d67338322" # 请替换为您的有效API Key
QIANWEN_MODEL = 'qwen-turbo'
MAX_CHARS_PER_TRANSLATE_CHUNK = 1800
FALLBACK_CHINESE_FONT = 'SimSun'


# --- Qianwen Translation Function ---
def translate_text_qianwen(text_to_translate: str, target_language: str = "Simplified Chinese") -> Union[str, None]:
    if not text_to_translate.strip(): return ""
    prompt = f"You are an expert translator specializing in avionics and telecommunications. Translate the following English text into {target_language}. Your translation must be precise, formal, and maintain the technical accuracy and tone of the original source. Do not add any extra explanations or introductory phrases like 'Here is the translation:'. Just provide the translated text directly.\n\nEnglish text to translate:\n\"\"\"\n{text_to_translate}\n\"\"\""
    try:
        if not dashscope.api_key or not dashscope.api_key.startswith("sk-"):
            print("Error: Qianwen API Key is invalid or not set for Dashscope.")
            return f"[Translation API Key Error]"
        response = dashscope.Generation.call(model=QIANWEN_MODEL, prompt=prompt)
        if response.status_code == HTTPStatus.OK:
            return response.output.text.strip() if response.output and response.output.text else f"[Translation Error: Empty Response]"
        else:
            print(
                f"Error calling Qianwen API: HTTP Status {response.status_code}, Code: {response.code}, Message: {response.message}")
            return f"[Translation API Error: {response.message}]"
    except Exception as e:
        print(f"Exception during Qianwen API call: {e}")
        return f"[Translation Exception: {e}]"


# --- Text Chunking Function ---
def split_text_into_chunks(long_text: str, max_chars: int) -> list[str]:
    if len(long_text) <= max_chars: return [long_text] if long_text.strip() else []
    try:
        sentences = nltk.sent_tokenize(long_text)
    except Exception as e:
        print(f"NLTK tokenization failed: {e}. Fallback to char split.")
        return [long_text[i:i + max_chars] for i in range(0, len(long_text), max_chars)]
    chunks, current_chunk = [], ""
    for sentence in sentences:
        if len(current_chunk) + len(sentence) + 1 > max_chars and current_chunk:
            chunks.append(current_chunk.strip());
            current_chunk = ""
        if len(sentence) > max_chars:
            if current_chunk: chunks.append(current_chunk.strip()); current_chunk = ""
            for i in range(0, len(sentence), max_chars): chunks.append(sentence[i:i + max_chars])
        else:
            current_chunk += (sentence + " ")
    if current_chunk: chunks.append(current_chunk.strip())
    return [c for c in chunks if c.strip()]

# --- DOCX Parsing Helper Functions ---
def _get_on_off_val(oxml_element, attribute_qname_val=qn('w:val')):
    if oxml_element is None:
        return False
    val = oxml_element.get(attribute_qname_val) 
    if val is None: 
        return True
    if isinstance(val, str) and (val.lower() == 'false' or val == '0'):
        return False
    if isinstance(val, bool) and not val: 
        return False
    return True

def get_run_formatting(r_element):
    style_info = {'bold': None, 'italic': None, 'underline': None, 'font_name': None, 'font_size': None,
                  'font_color_rgb': None, 'char_style_name': None}
    
    rPr = r_element.rPr
    if rPr is None:
        return style_info

    # Bold
    try:
        if ((rPr.b is not None and rPr.b.val is not False) or
            (rPr.bCs is not None and rPr.bCs.val is not False)): # Wrapped in outer parentheses
            style_info['bold'] = True
    except AttributeError: 
        if (_get_on_off_val(rPr.find(qn('w:b'))) or
            _get_on_off_val(rPr.find(qn('w:bCs')))): # Wrapped in outer parentheses
            style_info['bold'] = True

    # Italic
    try:
        if ((rPr.i is not None and rPr.i.val is not False) or
            (rPr.iCs is not None and rPr.iCs.val is not False)): # Wrapped in outer parentheses
            style_info['italic'] = True
    except AttributeError:
        if (_get_on_off_val(rPr.find(qn('w:i'))) or
            _get_on_off_val(rPr.find(qn('w:iCs')))): # Wrapped in outer parentheses
            style_info['italic'] = True
            
    # Underline
    try:
        u_obj = rPr.u
        if u_obj is not None and u_obj.val is not None and str(u_obj.val).lower() != 'none':
            style_info['underline'] = True
    except AttributeError:
        u_oxml_el = rPr.find(qn('w:u'))
        if u_oxml_el is not None:
            u_val_str = u_oxml_el.get(qn('w:val'))
            if u_val_str is not None and u_val_str.lower() != 'none':
                style_info['underline'] = True
                
    # Font Name
    font_name_val = None
    try:
        rFonts_obj = rPr.rFonts
        if rFonts_obj is not None:
            # Using parentheses for multi-line 'or' chain for clarity, \ not strictly needed here
            font_name_val = (getattr(rFonts_obj, 'eastAsia', None) or
                             getattr(rFonts_obj, 'ascii', None) or
                             getattr(rFonts_obj, 'hAnsi', None) or
                             getattr(rFonts_obj, 'cs', None))
            if not font_name_val: 
                font_name_val = (rFonts_obj.get(qn('w:eastAsia')) or
                                 rFonts_obj.get(qn('w:ascii')) or
                                 rFonts_obj.get(qn('w:hAnsi')) or
                                 rFonts_obj.get(qn('w:cs')))
    except AttributeError: 
        rFonts_oxml_el = rPr.find(qn('w:rFonts'))
        if rFonts_oxml_el is not None:
            font_name_val = (rFonts_oxml_el.get(qn('w:eastAsia')) or
                             rFonts_oxml_el.get(qn('w:ascii')) or
                             rFonts_oxml_el.get(qn('w:hAnsi')) or
                             rFonts_oxml_el.get(qn('w:cs')))
    if font_name_val:
        style_info['font_name'] = font_name_val

    # Font Size
    size_attribute_val = None
    try:
        sz_obj = rPr.sz
        szCs_obj = rPr.szCs 
        
        actual_sz_val_provider = szCs_obj if szCs_obj is not None and szCs_obj.val is not None else sz_obj
        if actual_sz_val_provider is not None and actual_sz_val_provider.val is not None:
            size_attribute_val = actual_sz_val_provider.val
    except AttributeError: 
        szCs_oxml_el = rPr.find(qn('w:szCs'))
        if szCs_oxml_el is not None and szCs_oxml_el.get(qn('w:val')) is not None:
            size_attribute_val = szCs_oxml_el.get(qn('w:val'))
        else:
            sz_oxml_el = rPr.find(qn('w:sz'))
            if sz_oxml_el is not None and sz_oxml_el.get(qn('w:val')) is not None:
                size_attribute_val = sz_oxml_el.get(qn('w:val'))
    
    if size_attribute_val is not None:
        try:
            style_info['font_size'] = Pt(int(size_attribute_val) / 2)
        except (ValueError, TypeError): pass

    # Font Color
    color_val_str = None
    try:
        color_obj = rPr.color
        if color_obj is not None and color_obj.val is not None:
            color_val_str = str(color_obj.val)
    except AttributeError:
        color_oxml_el = rPr.find(qn('w:color'))
        if color_oxml_el is not None and color_oxml_el.get(qn('w:val')) is not None:
            color_val_str = str(color_oxml_el.get(qn('w:val')))
            
    if color_val_str and color_val_str.lower() != 'auto':
        style_info['font_color_rgb'] = color_val_str

    # Character Style
    char_style_val = None
    try:
        rStyle_obj = rPr.rStyle
        if rStyle_obj is not None and rStyle_obj.val is not None:
            char_style_val = rStyle_obj.val
    except AttributeError:
        rStyle_oxml_el = rPr.find(qn('w:rStyle'))
        if rStyle_oxml_el is not None and rStyle_oxml_el.get(qn('w:val')) is not None:
            char_style_val = rStyle_oxml_el.get(qn('w:val'))
    if char_style_val:
        style_info['char_style_name'] = char_style_val
            
    return style_info

def convert_image_to_png_bytes(image_bytes: bytes, original_filename: str) -> Union[bytes, None]:
    """
    尝试将传入的图像字节流（假定为EMF/WMF）转换为PNG字节流。
    使用外部命令行工具。
    返回PNG字节流，如果转换失败则返回None。
    """
    _, original_extension = os.path.splitext(original_filename)
    png_bytes = None

    # 为了更好地管理临时文件名并避免冲突，我们可以让系统命名它们
    tmp_in_fd, tmp_in_path = tempfile.mkstemp(suffix=original_extension)
    tmp_out_fd, tmp_out_path = tempfile.mkstemp(suffix=".png")

    # 首先关闭文件描述符，因为我们只是用它们来获取安全的临时文件名
    # 稍后将以二进制模式重新打开这些文件进行读写
    os.close(tmp_in_fd)
    os.close(tmp_out_fd)

    try:
        with open(tmp_in_path, 'wb') as f_in:
            f_in.write(image_bytes)

        # --- 使用 ImageMagick v7+ 的 'magick' 命令 ---
        cmd = ['magick', tmp_in_path, tmp_out_path] 
        # 或者更明确的： cmd = ['magick', 'convert', tmp_in_path, tmp_out_path]

        # ---- 以下是备选方案，如果ImageMagick+LibreOffice方式持续有问题 ----
        # 备选 1: 如果您安装并希望直接使用 unoconv (需要 LibreOffice 和 unoconv)
        # 您需要先安装 unoconv: brew install unoconv (或者 pip install unoconv)
        # cmd = ['unoconv', '-f', 'png', '-o', tmp_out_path, tmp_in_path]

        # 备选 2: 如果您安装并希望直接使用 Inkscape (需要 Inkscape)
        # inkscape_path = '/Applications/Inkscape.app/Contents/MacOS/inkscape' # macOS上Inkscape的典型路径
        # if not os.path.exists(inkscape_path): # 检查inkscape是否存在
        #     inkscape_path = 'inkscape' # 如果在PATH中，直接用inkscape
        # cmd = [inkscape_path, f'--export-filename={tmp_out_path}', '--export-type=png', tmp_in_path]
        # ---- 备选方案结束 ----

        print(f"执行转换命令: {' '.join(cmd)}")
        process = subprocess.run(cmd, capture_output=True, timeout=60, check=False)

        if process.returncode == 0:
            if os.path.exists(tmp_out_path) and os.path.getsize(tmp_out_path) > 0:
                with open(tmp_out_path, 'rb') as f_png:
                    png_bytes = f_png.read()
            else:
                 print(f"警告: 图像转换命令 '{' '.join(cmd)}' 可能已执行，但输出PNG文件 '{tmp_out_path}' 未找到或为空。")
                 if process.stderr: print(f"Stderr: {process.stderr.decode(errors='ignore')}")
                 if process.stdout: print(f"Stdout: {process.stdout.decode(errors='ignore')}")
        else:
            print(f"警告: '{original_filename}' 到 PNG 的转换失败。返回码: {process.returncode}")
            print(f"命令: {' '.join(cmd)}")
            if process.stderr: print(f"Stderr: {process.stderr.decode(errors='ignore')}")
            if process.stdout: print(f"Stdout: {process.stdout.decode(errors='ignore')}")

    except FileNotFoundError:
        tool_name = cmd[0] if 'cmd' in locals() and cmd else "转换工具"
        print(f"错误: {tool_name} 未找到。请确保已安装该工具并将其添加至系统PATH。")
    except subprocess.TimeoutExpired:
        print(f"错误: '{original_filename}' 到 PNG 的转换超时。")
    except Exception as e:
        print(f"'{original_filename}' 到 PNG 的转换过程中发生意外错误: {e}")
    finally:
        if os.path.exists(tmp_in_path):
            os.remove(tmp_in_path)
        if os.path.exists(tmp_out_path):
            os.remove(tmp_out_path)

    return png_bytes


def extract_image_details_from_drawing(drawing_element, doc_part):
    try:
        embed_rId = None
        if drawing_element.tag == qn('w:drawing'):
            blip_element = drawing_element.find('.//' + qn('a:blip'))
            if blip_element is not None:
                embed_rId = blip_element.get(qn('r:embed'))
        elif drawing_element.tag == qn('v:imagedata'): 
            embed_rId = drawing_element.get(qn('r:id'))

        if not embed_rId: return None

        image_part = doc_part.related_parts[embed_rId]
        image_blob, image_filename = image_part.blob, os.path.basename(image_part.partname)
        
        extent_element = None
        width, height = None, None

        if drawing_element.tag == qn('w:drawing'):
             extent_element = drawing_element.find('.//' + qn('wp:extent')) 
             if extent_element is None:
                graphic_data_extent = drawing_element.find(
                    './/' + qn('a:graphicData') + '/' + qn('pic:pic') + '/' + qn('pic:spPr') + '/' + qn(
                        'a:xfrm') + '/' + qn('a:ext'))
                if graphic_data_extent is not None: extent_element = graphic_data_extent
        
        if extent_element is not None:
            cx_str, cy_str = extent_element.get('cx'), extent_element.get('cy')
            if cx_str and cy_str:
                try:
                    width, height = Emu(int(cx_str)), Emu(int(cy_str))
                except ValueError:
                    pass
        return {'stream': io.BytesIO(image_blob), 'width': width, 'height': height, 'filename': image_filename}
    except KeyError: 
        rId_val = embed_rId if 'embed_rId' in locals() and embed_rId is not None else 'unknown_rId'; 
        print(f"Error: rId '{rId_val}' not found in document part relationships. Skipping image."); return None
    except Exception as e:
        print(f"Error extracting image details for element {drawing_element.tag if hasattr(drawing_element, 'tag') else 'unknown'}: {e}"); return None


def _parse_single_paragraph_content(p_obj: DocxParagraphObject, source_doc_part):
    content_items = []
    text_for_translation_parts = []
    for r_element in p_obj._p.xpath('./w:r'): 
        current_run_text = "".join([t.text for t in r_element.xpath('./w:t') if t.text is not None])
        run_style_info = get_run_formatting(r_element) 
        
        image_drawing_element = None
        std_drawing_elements = r_element.xpath('./w:drawing')
        if std_drawing_elements:
            image_drawing_element = std_drawing_elements[0]
        else:
            pict_elements = r_element.xpath('./w:pict') 
            for pict_element in pict_elements: 
                try:
                    vml_imagedata_elements = pict_element.xpath('.//v:shape//v:imagedata', namespaces=VML_NS_MAP)
                    if vml_imagedata_elements:
                        image_drawing_element = vml_imagedata_elements[0] 
                        break 
                except Exception: pass 

        if image_drawing_element is not None:
            image_details = extract_image_details_from_drawing(image_drawing_element, source_doc_part)
            if image_details:
                # --- 修改后的图像处理逻辑 ---
                original_filename = image_details.get('filename', '')
                if original_filename and \
                        (original_filename.lower().endswith('.emf') or original_filename.lower().endswith('.wmf')):
                    
                    print(f"检测到 EMF/WMF 图片: '{original_filename}'. 将直接嵌入原始图像。")
                    image_details['stream'].seek(0) # 确保流指针在开头以便后续读取
                    # 由于我们直接嵌入原始EMF/WMF，保留其原始文件名和提取到的尺寸信息
                    # image_details['filename'] 保持不变
                    # image_details['width'] 和 image_details['height'] 保持从 extract_image_details_from_drawing 获取的值
                else:
                    # 对于非 EMF/WMF 图片 (例如 PNG, JPEG 等)
                    # 如果它们也需要转换或特殊处理，可以在这里添加逻辑。
                    # 目前脚本的行为是对它们使用原始流，这通常是期望的。
                    # 如果它们是脚本中其他地方转换失败的产物，也确保流指针在开头。
                    print(f"处理其他格式图片: '{original_filename}'. 将使用原始图像数据。")
                    image_details['stream'].seek(0)
                
                content_items.append({'type': 'image', **image_details})
                # --- 图像处理逻辑结束 ---
            
            # 如果图片伴随文本 (虽然不常见于同一run内同时有图有字，但以防万一)
            if current_run_text:
                content_items.append({'type': 'text', 'text': current_run_text, 'style_info': run_style_info})
                text_for_translation_parts.append(current_run_text)
        elif current_run_text: # 只有文本
            content_items.append({'type': 'text', 'text': current_run_text, 'style_info': run_style_info})
            text_for_translation_parts.append(current_run_text)
            
    return {
        'content_items': content_items,
        'text_to_translate': "".join(text_for_translation_parts),
        'paragraph_style_name': p_obj.style.name if p_obj.style and p_obj.style.name else None,
        'paragraph_alignment': p_obj.alignment
    }


def parse_docx_for_translation(source_doc: Document) -> list:
    parsed_document_structure = []
    for block_element in source_doc.element.body.xpath('./w:p[not(ancestor::w:tbl)]|./w:tbl'):
        if block_element.tag == qn('w:p'):
            para_obj = DocxParagraphObject(block_element, source_doc)
            parsed_para_data = _parse_single_paragraph_content(para_obj, source_doc.part)
            current_paragraph_data = {'type': 'paragraph', **parsed_para_data, 'translation': None}
            if current_paragraph_data['content_items'] or current_paragraph_data['text_to_translate'].strip():
                parsed_document_structure.append(current_paragraph_data)

        elif block_element.tag == qn('w:tbl'):
            table_obj = DocxTableObject(block_element, source_doc)
            table_data = {'type': 'table', 'rows': [],
                          'style_name': table_obj.style.name if table_obj.style and table_obj.style.name else None,
                          'max_cols': len(table_obj.columns) if table_obj.columns else 0}

            max_cols_actual = 0
            for r_idx, row_obj in enumerate(table_obj.rows):
                row_data = {'cells': []}
                current_col_span_in_row = 0
                for c_idx, cell_obj in enumerate(row_obj.cells):
                    cell_paragraphs_data_list, cell_text_parts_for_translation = [], []
                    for p_in_cell_obj in cell_obj.paragraphs:
                        parsed_para_in_cell = _parse_single_paragraph_content(p_in_cell_obj, source_doc.part)
                        if parsed_para_in_cell['content_items'] or parsed_para_in_cell['text_to_translate'].strip():
                            cell_paragraphs_data_list.append(parsed_para_in_cell)
                            cell_text_parts_for_translation.append(parsed_para_in_cell['text_to_translate'])
                        elif not cell_paragraphs_data_list and not cell_obj.text.strip():
                             cell_paragraphs_data_list.append(parsed_para_in_cell)

                    grid_span_val, v_merge_val = 1, None
                    tcPr = cell_obj._tc.tcPr
                    if tcPr is not None:
                        gs_el = tcPr.gridSpan
                        if gs_el is not None: grid_span_val = gs_el.val if gs_el.val is not None else 1
                        
                        vm_el = tcPr.vMerge
                        if vm_el is not None:
                            v_merge_val = vm_el.val 
                            if v_merge_val is None: 
                                v_merge_val = 'continue'
                                
                    cell_data = {'paragraphs_data': cell_paragraphs_data_list,
                                 'full_cell_text_to_translate': "".join(cell_text_parts_for_translation),
                                 'translation': None, 'grid_span': grid_span_val, 'v_merge': v_merge_val}
                    row_data['cells'].append(cell_data)
                    current_col_span_in_row += grid_span_val
                table_data['rows'].append(row_data)
                if current_col_span_in_row > max_cols_actual: max_cols_actual = current_col_span_in_row
            
            table_data['max_cols'] = max_cols_actual if max_cols_actual > 0 else table_data['max_cols']
            parsed_document_structure.append(table_data)
    return parsed_document_structure


# --- DOCX Writing Helper Functions ---
def apply_run_formatting(target_run, style_info: dict):
    if style_info.get('bold') is True: target_run.bold = True
    if style_info.get('italic') is True: target_run.italic = True
    if style_info.get('underline') is True: target_run.underline = True
    if style_info.get('font_name'): target_run.font.name = style_info['font_name']
    if style_info.get('font_size'): target_run.font.size = style_info['font_size']
    hex_color = style_info.get('font_color_rgb')
    if hex_color and hex_color.lower() != 'auto' and hex_color.lower() != "000000":
        try:
            if len(hex_color) == 6: target_run.font.color.rgb = RGBColor.from_string(hex_color)
        except Exception:
            pass
    char_style_name = style_info.get('char_style_name')
    if char_style_name and char_style_name in target_run.part.document.styles:
        try:
            target_run.style = char_style_name
        except KeyError:
            pass

def create_interleaved_docx(output_file_path: str, processed_data: list, source_doc: Document):
    new_doc = Document()

    for data_item in tqdm(processed_data, desc="生成对照文档", unit="元素"):
        if data_item['type'] == 'paragraph':
            eng_para_style_name = data_item.get('paragraph_style_name')
            try:
                eng_p = new_doc.add_paragraph(style=eng_para_style_name if eng_para_style_name and eng_para_style_name in new_doc.styles else None)
            except KeyError: 
                eng_p = new_doc.add_paragraph()
            if data_item.get('paragraph_alignment') is not None: eng_p.alignment = data_item['paragraph_alignment']

            for item in data_item['content_items']:
                if item['type'] == 'text':
                    run = eng_p.add_run(item['text'])
                    apply_run_formatting(run, item['style_info'])
                elif item['type'] == 'image':
                    try:
                        item['stream'].seek(0)
                        pic_run = eng_p.add_run()
                        if item['width'] and item['height']:
                            pic_run.add_picture(item['stream'], width=item['width'], height=item['height'])
                        elif item['width']:
                            pic_run.add_picture(item['stream'], width=item['width'])
                        else:
                            pic_run.add_picture(item['stream']) 
                    except Exception as e:
                        print(f"Error adding image '{item.get('filename', 'unknown')}' to paragraph: {e}")

            translated_text = data_item.get('translation')
            if translated_text and translated_text.strip():
                zh_para_style_name = data_item.get('paragraph_style_name')
                try:
                    zh_p = new_doc.add_paragraph(style=zh_para_style_name if zh_para_style_name and zh_para_style_name in new_doc.styles else None)
                except KeyError: 
                    zh_p = new_doc.add_paragraph()
                if data_item.get('paragraph_alignment') is not None: zh_p.alignment = data_item['paragraph_alignment']
                zh_run = zh_p.add_run(translated_text)
                zh_run.font.name = FALLBACK_CHINESE_FONT
                try:
                    rpr = zh_run._r.get_or_add_rPr()
                    rfonts = rpr.get_or_add_rFonts()
                    rfonts.set(qn('w:eastAsia'), FALLBACK_CHINESE_FONT)
                except Exception: pass

        elif data_item['type'] == 'table':
            num_rows = len(data_item['rows'])
            num_cols = data_item.get('max_cols', 1)
            if num_rows == 0 or num_cols == 0: continue

            new_table = new_doc.add_table(rows=num_rows, cols=num_cols)
            table_style_name = data_item.get('style_name')

            if table_style_name and table_style_name in new_doc.styles:
                try: new_table.style = table_style_name
                except Exception as e_style_apply:
                    print(f"警告: 应用表格样式 '{table_style_name}' 失败。错误: {e_style_apply}")
                    if 'TableGrid' in new_doc.styles: new_table.style = 'TableGrid'
            elif table_style_name:
                if 'TableGrid' in new_doc.styles: new_table.style = 'TableGrid'
                print(f"警告: 表格样式 '{table_style_name}' 在目标文档中未找到，已使用 'TableGrid'。")
            else:
                if 'TableGrid' in new_doc.styles: new_table.style = 'TableGrid'

            for r_idx, row_data in enumerate(data_item['rows']):
                current_col_idx_in_new_table = 0
                for c_idx_orig, cell_data in enumerate(row_data['cells']):
                    if current_col_idx_in_new_table >= num_cols: break 

                    new_cell = new_table.cell(r_idx, current_col_idx_in_new_table)
                    if new_cell.paragraphs and len(new_cell.paragraphs) == 1 and not new_cell.paragraphs[0].text and not new_cell.paragraphs[0].runs:
                        p_elm = new_cell.paragraphs[0]._element
                        if p_elm.getparent() is not None : p_elm.getparent().remove(p_elm)
                    
                    for p_struct_idx, p_struct_in_cell in enumerate(cell_data['paragraphs_data']):
                        cell_eng_p = new_cell.add_paragraph()

                        cell_eng_p_style = p_struct_in_cell.get('paragraph_style_name')
                        if cell_eng_p_style and cell_eng_p_style in new_doc.styles:
                            try: cell_eng_p.style = cell_eng_p_style
                            except KeyError: pass 

                        cell_eng_p_align = p_struct_in_cell.get('paragraph_alignment')
                        if cell_eng_p_align is not None: cell_eng_p.alignment = cell_eng_p_align

                        for item in p_struct_in_cell['content_items']:
                            if item['type'] == 'text':
                                run = cell_eng_p.add_run(item['text'])
                                apply_run_formatting(run, item['style_info'])
                            elif item['type'] == 'image':
                                try:
                                    item['stream'].seek(0)
                                    pic_run = cell_eng_p.add_run()
                                    if item['width'] and item['height']: pic_run.add_picture(item['stream'], width=item['width'], height=item['height'])
                                    elif item['width']: pic_run.add_picture(item['stream'], width=item['width'])
                                    else: pic_run.add_picture(item['stream'])
                                except Exception as e:
                                    print(f"Error adding image '{item.get('filename', 'unknown')}' to table cell: {e}")
                    
                    cell_translation = cell_data.get('translation')
                    if cell_translation and cell_translation.strip():
                        cell_zh_p = new_cell.add_paragraph()
                        zh_run = cell_zh_p.add_run(cell_translation)
                        zh_run.font.name = FALLBACK_CHINESE_FONT
                        try:
                            rpr = zh_run._r.get_or_add_rPr()
                            rfonts = rpr.get_or_add_rFonts()
                            rfonts.set(qn('w:eastAsia'), FALLBACK_CHINESE_FONT)
                        except Exception: pass
                    
                    grid_span = cell_data.get('grid_span', 1)
                    if grid_span > 1:
                        end_col_idx = current_col_idx_in_new_table + grid_span - 1
                        if end_col_idx < num_cols:
                            try:
                                new_cell.merge(new_table.cell(r_idx, end_col_idx))
                            except Exception as e_h_merge:
                                print(f"警告: 水平合并单元格 ({r_idx},{current_col_idx_in_new_table}) 到 ({r_idx},{end_col_idx}) 失败: {e_h_merge}")
                    current_col_idx_in_new_table += grid_span
            
            for r_idx_vmerge, row_data_vmerge in enumerate(data_item['rows']):
                current_col_idx_for_vmerge_setting = 0
                for source_cell_d_vmerge in row_data_vmerge['cells']:
                    v_merge_val_from_parser = source_cell_d_vmerge.get('v_merge')

                    if v_merge_val_from_parser is not None: 
                        if current_col_idx_for_vmerge_setting < num_cols :
                            cell_to_set_vmerge_on = new_table.cell(r_idx_vmerge, current_col_idx_for_vmerge_setting)
                            # Get or create the <w:vMerge> element. Its mere presence indicates a merge.
                            vmerge_child_element = cell_to_set_vmerge_on._tc.get_or_add_tcPr().get_or_add_vMerge()
                            # Only set the 'val' attribute if it's a 'restart'.
                            # For 'continue', the attribute should be absent.
                            if v_merge_val_from_parser == 'restart':
                                vmerge_child_element.val = 'restart'
                                 
                    current_col_idx_for_vmerge_setting += source_cell_d_vmerge.get('grid_span', 1)

    new_doc.save(output_file_path)
    print(f"交叉翻译文档已保存到: {output_file_path}")


# --- Main Orchestration Function ---
def translate_document_main_flow(input_docx_path: str, output_docx_path: str):
    print(f"开始翻译流程: {input_docx_path} -> {output_docx_path}")
    try:
        source_doc_obj = Document(input_docx_path)
    except Exception as e:
        print(f"错误：无法打开源文档 '{input_docx_path}'. Error: {e}"); return False

    parsed_structure = parse_docx_for_translation(source_doc_obj)
    if not parsed_structure: print("解析文档结构失败或文档为空。"); return False

    print(f"成功解析 {len(parsed_structure)} 个顶层元素。开始收集待翻译内容...")

    # 1. 收集所有待翻译的任务
    # 每个任务是一个字典，包含待翻译文本和结果应存回的目标字典的引用。
    tasks_to_translate = []
    for item in parsed_structure:
        if item['type'] == 'paragraph' and item.get('text_to_translate', '').strip():
            tasks_to_translate.append({'text': item['text_to_translate'], 'target': item})
        elif item['type'] == 'table':
            for row in item['rows']:
                for cell in row['cells']:
                    if cell.get('full_cell_text_to_translate', '').strip():
                        tasks_to_translate.append({'text': cell['full_cell_text_to_translate'], 'target': cell})

    if not tasks_to_translate:
        print("未找到需要翻译的文本内容。如果文档中确实有文本，请检查解析逻辑。")
        create_interleaved_docx(output_docx_path, parsed_structure, source_doc_obj)
        return True

    # 2. 定义一个辅助函数来处理单个文本块的翻译（包括分片）
    def process_and_translate_block(text_block: str) -> str:
        chunks = split_text_into_chunks(text_block, MAX_CHARS_PER_TRANSLATE_CHUNK)
        translated_chunks = []
        for chunk in chunks:
            translated_chunk = translate_text_qianwen(chunk)
            if translated_chunk and not translated_chunk.startswith("[Translation"):
                translated_chunks.append(translated_chunk)
            else:
                print(f"警告: 一块文本翻译失败或返回空: '{chunk[:50]}...'。错误信息: {translated_chunk}")
                translated_chunks.append(f"[翻译失败: {chunk[:30]}...]")
        return " ".join(filter(None, translated_chunks)).strip()

    # 3. 使用线程池并发执行翻译
    # max_workers可以根据您的API速率限制和网络状况调整，10是一个不错的起点。
    with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
        texts_for_api = [task['text'] for task in tasks_to_translate]
        
        # executor.map会按顺序返回结果，非常适合与tqdm结合使用来显示进度
        results_iterator = executor.map(process_and_translate_block, texts_for_api)
        translated_texts = list(tqdm(results_iterator, total=len(tasks_to_translate), desc="翻译进度", unit="单元"))

    # 4. 将翻译结果写回数据结构
    for i, task in enumerate(tasks_to_translate):
        task['target']['translation'] = translated_texts[i]

    print("文本翻译阶段完成。")
    create_interleaved_docx(output_docx_path, parsed_structure, source_doc_obj)
    return True


# --- 主程序入口 ---
if __name__ == '__main__':
    if not QIANWEN_API_KEY or not QIANWEN_API_KEY.startswith("sk-") or len(QIANWEN_API_KEY) < 20:
        print("错误：千问API Key (QIANWEN_API_KEY) 无效或未在脚本中正确设置。请确保API Key以'sk-'开头且长度合理。")
        sys.exit(1)
    else:
        dashscope.api_key = QIANWEN_API_KEY
        print(f"千问 API Key 已设置 (Key ending with: ...{QIANWEN_API_KEY[-4:]}).")

        if len(sys.argv) < 2:
            script_name = os.path.basename(sys.argv[0])
            print(f"使用方法: python3 {script_name} <待翻译的docx文件路径>")
            sys.exit(1) 

        input_file_path = sys.argv[1]

        if not os.path.exists(input_file_path):
            print(f"错误: 输入文件 '{input_file_path}' 未找到。")
            sys.exit(1) 
        
        if not input_file_path.lower().endswith(".docx"):
            print(f"错误: 输入文件 '{input_file_path}' 不是一个 .docx 文件。")
            sys.exit(1)

        input_dir = os.path.dirname(os.path.abspath(input_file_path))
        input_filename_only = os.path.basename(input_file_path)
        
        base, ext = os.path.splitext(input_filename_only)
        output_filename = f"{base}_zh{ext}"
        output_file_path = os.path.join(input_dir, output_filename)

        try:
            print(f"输入文件: {input_file_path}")
            print(f"输出文件: {output_file_path}")
            translate_document_main_flow(input_file_path, output_file_path)

        except Exception as e:
            print(f"主程序执行过程中发生全局性意外错误: {e}")
            import traceback
            traceback.print_exc()